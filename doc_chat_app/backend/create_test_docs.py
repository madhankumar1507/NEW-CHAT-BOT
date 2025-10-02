
import docx
import os

# Define the directory to save the files
DOCS_DIR = os.path.join(os.path.dirname(__file__), "data", "docs")
os.makedirs(DOCS_DIR, exist_ok=True)

# --- Content for File 1: Overview ---
overview_text = [
    ("Mahabharata: An Overview", "h1"),
    ("The Mahabharata is one of the two major Sanskrit epics of ancient India, the other being the Ramayana. It is an epic narrative of the Kurukshetra War and the fates of the Kaurava and the Pandava princes. It also contains philosophical and devotional material, such as a discussion of the four 'goals of life' or purusharthas.", "p"),
    ("Traditionally, the authorship of the Mahabharata is attributed to Vyasa. With about 1.8 million words in total, the Mahabharata is roughly ten times the length of the Iliad and the Odyssey combined, or about four times the length of the Ramayana.", "p"),
    ("The story culminates in the great battle of Kurukshetra, which lasted for eighteen days, where the Pandavas are ultimately victorious. The Mahabharata ends with the death of Krishna, and the end of his dynasty and ascent of the Pandava brothers to heaven. It also marks the beginning of the Hindu age of Kali Yuga, the fourth and final age of humanity.", "p")
]

# --- Content for File 2: Characters ---
characters_text = [
    ("Key Characters in the Mahabharata", "h1"),
    ("The Pandavas", "h2"),
    ("The five sons of Pandu: Yudhishthira, Bhima, Arjuna, Nakula, and Sahadeva. They are the central protagonists of the epic.", "p"),
    ("The Kauravas", "h2"),
    ("The hundred sons of the blind king Dhritarashtra and his queen Gandhari. Duryodhana is the eldest and most prominent among them, acting as the primary antagonist.", "p"),
    ("Krishna", "h2"),
    ("A divine incarnation (avatar) of Vishnu, Krishna is a key figure who serves as a guide and charioteer for Arjuna. His counsel is pivotal to the Pandavas' efforts.", "p"),
    ("Draupadi", "h2"),
    ("The wife of the five Pandava brothers. Her humiliation at the hands of the Kauravas is a critical event that leads to the war.", "p")
]

# --- Content for File 3: Bhagavad Gita ---
gita_text = [
    ("The Bhagavad Gita", "h1"),
    ("The Bhagavad Gita, often referred to as the Gita, is a 700-verse Hindu scripture that is part of the epic Mahabharata. It is set in a narrative framework of a dialogue between Pandava prince Arjuna and his guide and charioteer Krishna.", "p"),
    ("At the start of the Kurukshetra War, Arjuna is filled with moral dilemma and despair about the violence and death the war will cause. He turns to Krishna for counsel, and the Gita is the record of their conversation.", "p"),
    ("Krishna imparts to Arjuna the wisdom of selfless action (Karma Yoga), the nature of the self (Atman), and the path to liberation (Moksha). The Gita's teachings have had a profound influence on Hindu thought and philosophy for centuries.", "p")
]

def create_doc(filename, content):
    doc = docx.Document()
    for text, style in content:
        if style == "h1":
            doc.add_heading(text, level=1)
        elif style == "h2":
            doc.add_heading(text, level=2)
        else:
            doc.add_paragraph(text)
    
    filepath = os.path.join(DOCS_DIR, filename)
    doc.save(filepath)
    print(f"Successfully created {filepath}")

# Create the documents
create_doc("mahabharata_overview.docx", overview_text)
create_doc("mahabharata_characters.docx", characters_text)
create_doc("bhagavad_gita.docx", gita_text)

print("\nTest files have been created in the 'backend/data/docs' directory.")
